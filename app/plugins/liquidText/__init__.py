import uuid
from app.utils.PluginClass import PluginClass
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.utils import DatabaseHandler
from flask import request, send_file
from flask_babel import _
from celery import shared_task
from dotenv import load_dotenv
import os
from app.api.resources.models import ResourceUpdate
from bson.objectid import ObjectId
import pandas as pd
import datetime
import shutil

load_dotenv()

mongodb = DatabaseHandler.DatabaseHandler()

USER_FILES_PATH = os.environ.get('USER_FILES_PATH', '')
WEB_FILES_PATH = os.environ.get('WEB_FILES_PATH', '')
ORIGINAL_FILES_PATH = os.environ.get('ORIGINAL_FILES_PATH', '')
TEMPORAL_FILES_PATH = os.environ.get('TEMPORAL_FILES_PATH', '')

class ExtendedPluginClass(PluginClass):
    def __init__(self, path, import_name, name, description, version, author, type, settings, actions, capabilities=None):
        super().__init__(path, __file__, import_name, name, description, version, author, type, settings, actions = actions, capabilities=None)

    def add_routes(self):
        @self.route('/bulk', methods=['POST'])
        @jwt_required()
        def generate_liquid_text():
            current_user = get_jwt_identity()
            body = request.get_json()
            
            if not self.has_role('admin', current_user) and not self.has_role('processing', current_user):
                return {'msg': 'No tiene permisos suficientes'}, 401
            
            from app.utils.functions import cache_get_processing_result
            cache_get_processing_result.invalidate_all()

            task = self.bulk.delay(body, current_user)
            self.add_task_to_user(task.id, 'liquidText.generateLiquidText', current_user, 'msg')
            
            return {'msg': 'Se agregó la tarea a la fila de procesamientos'}, 201
        
        @self.route('/filedownload/<taskId>', methods=['GET'])
        @jwt_required()
        def file_download(taskId):
            current_user = get_jwt_identity()

            if not self.has_role('admin', current_user) and not self.has_role('processing', current_user) and not self.has_role('editor', current_user):
                return {'msg': 'No tiene permisos suficientes'}, 401
            
            # Buscar la tarea en la base de datos
            task = mongodb.get_record('tasks', {'taskId': taskId})
            # Si la tarea no existe, retornar error
            if not task:
                return {'msg': 'Tarea no existe'}, 404
            
            if task['user'] != current_user and not self.has_role('admin', current_user):
                return {'msg': 'No tiene permisos para obtener la tarea'}, 401

            if task['status'] == 'pending':
                return {'msg': 'Tarea en proceso'}, 400

            if task['status'] == 'failed':
                return {'msg': 'Tarea fallida'}, 400

            if task['status'] == 'completed':
                if task['resultType'] != 'file_download':
                    return {'msg': 'Tarea no es de tipo file_download'}, 400
                
            path = USER_FILES_PATH + task['result']
            response = send_file(path, as_attachment=True, download_name=os.path.basename(path), conditional=False)
            
            response.headers.add("Access-Control-Expose-Headers", "Content-Disposition")
            return response
        
        @self.route('/download', methods=['POST'])
        @jwt_required()
        def download_liquid_text():
            current_user = get_jwt_identity()
            body = request.get_json()
            
            if not self.has_role('admin', current_user) and not self.has_role('processing', current_user):
                return {'msg': 'No tiene permisos suficientes'}, 401

            task = self.download.delay(body, current_user)
            self.add_task_to_user(task.id, 'liquidText.downloadLiquidText', current_user, 'file_download')
            
            return {'msg': 'Se agregó la tarea a la fila de procesamientos'}, 201
        
        @self.route('/save', methods=['POST'])
        @jwt_required()
        def save_liquid_text():
            current_user = get_jwt_identity()
            body = request.get_json()
            
            if 'id' not in body:
                return {'msg': 'No se ha enviado el id del recurso'}, 400
            
            if not self.has_role('admin', current_user) and not self.has_role('processing', current_user):
                return {'msg': 'No tiene permisos suficientes'}, 401

            resp = self.save(body, current_user)
            
            return {'msg': resp}, 201
        
    def save(self, body, user):
        instance = ExtendedPluginClass('liquidText', '', **plugin_info)
        
        record = mongodb.get_record('records', {'_id': ObjectId(body['id'])}, fields={'_id': 1, 'processing': 1, 'name': 1})
        if not record:
            return {'msg': 'No se ha encontrado el recurso'}, 404
        
        if 'processing' not in record:
            return {'msg': 'No se ha encontrado el recurso'}, 404
        
        if body['slug'] not in record['processing']:
            return {'msg': 'No se ha encontrado el recurso'}, 404
        
        update = {
            'processing': record['processing']
        }
        
        update['processing']['liquidText'] = {
            'type': 'liquidText',
            'result': {
                'text': body['text'],
                'status': 'completed',
                'date': datetime.datetime.now()
            }
        }
        instance.update_data('records', str(record['_id']), update)
        
        from app.utils.functions import cache_get_processing_result
        cache_get_processing_result.invalidate_all()
        
        resp = _("Se ha guardado el texto líquido en el recurso %(resource)s", resource=record['name'])
        return resp
        
    @shared_task(ignore_result=False, name='liquidText.downloadLiquidText')
    def download(body, user):
        records_filters = {'_id': {'$in': [ObjectId(record) for record in body['records']]}}
        
        records = list(mongodb.get_all_records('records', records_filters, fields={'_id': 1, 'mime': 1, 'filepath': 1, 'processing': 1, 'name': 1, 'displayName': 1}))
        
        if len(records) == 0:
            raise Exception('No se encontraron registros')
        elif len(records) > 1:
            raise Exception('Debe seleccionar solo un registro')
        
        folder_path = USER_FILES_PATH + '/' + user + '/liquidText'
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
            
        def doc_processing(result, path):
            
            from docx import Document
            from bs4 import BeautifulSoup
            import re

            doc = Document()
            title = r['displayName'] if 'displayName' in r else r['name']
            doc.add_heading(title, 0)

            soup = BeautifulSoup(result, 'html.parser')

            for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                if tag.name == 'p':
                    paragraph = doc.add_paragraph()
                    for content in tag.contents:
                        if content.name in ['b', 'strong']:
                            run = paragraph.add_run(content.get_text())
                            run.bold = True
                        elif content.name in ['i', 'em']:
                            run = paragraph.add_run(content.get_text())
                            run.italic = True
                        elif content.name == 'u':
                            run = paragraph.add_run(content.get_text())
                            run.underline = True
                        elif content.name is None:
                            paragraph.add_run(str(content))
                elif re.match(r'h[1-6]', tag.name):
                    level = int(tag.name[1])
                    heading = doc.add_heading(level=level)
                    for content in tag.contents:
                        if content.name in ['b', 'strong']:
                            run = heading.add_run(content.get_text())
                            run.bold = True
                        elif content.name in ['i', 'em']:
                            run = heading.add_run(content.get_text())
                            run.italic = True
                        elif content.name == 'u':
                            run = heading.add_run(content.get_text())
                            run.underline = True
                        elif content.name is None:
                            heading.add_run(str(content))
                    
            doc.save(path)
            return path
        
        for r in records:
            if 'liquidText' in r['processing']:
                result = r['processing']['liquidText']['result']['text']
                
                if body['format'] == 'doc':
                    path = os.path.join(folder_path, str(r['_id']) + '.docx')
                    doc_processing(result, path)
                    return '/' + user + '/liquidText/' + str(r['_id']) + '.docx'
                    
                elif body['format'] == 'pdf':
                    temp_path = os.path.join(TEMPORAL_FILES_PATH, str(r['_id']) + '.docx')
                    doc_processing(result, temp_path)
                    
                    try:
                        from app.plugins.filesProcessing.utils.DocumentProcessing import convert_to_pdf_with_libreoffice
                    except Exception as e:
                        raise Exception('Error al importar el módulo del plugin para el procesamiento de documentos: ' + str(e))
                    
                    output_pdf = os.path.join(USER_FILES_PATH, user, 'liquidText', str(r['_id']) + '.pdf')
                    convert_to_pdf_with_libreoffice(temp_path, output_pdf)
                    shutil.move(os.path.join(TEMPORAL_FILES_PATH, str(r['_id']) + '.pdf'), output_pdf)
                    os.remove(temp_path)
                    return '/' + user + '/liquidText/' + str(r['_id']) + '.pdf'

                elif body['format'] == 'txt':
                    path = os.path.join(folder_path, str(r['_id']) + '.txt')
                    with open(path, 'w') as f:
                        f.write(result)
                    return '/' + user + '/liquidText/' + str(r['_id']) + '.txt'
                else:
                    raise Exception('Formato no soportado')

        return 'ok'
        
    @shared_task(ignore_result=False, name='liquidText.generateLiquidText')
    def bulk(body, user):
        instance = ExtendedPluginClass('liquidText', '', **plugin_info)
        if 'records' not in body:
            filters = {
                'post_type': body['post_type']
            }
            
            if isinstance(body['post_type'], list):
                filters['post_type'] = {'$in': body['post_type']}   

            if 'parent' in body:
                if body['parent'] and len(body['resources']) == 0:
                    filters = {'$or': [{'parents.id': body['parent'], 'post_type': filters['post_type']}, {'_id': ObjectId(body['parent'])}]}
            
            if 'resources' in body:
                if body['resources']:
                    if len(body['resources']) > 0:
                        filters = {'_id': {'$in': [ObjectId(resource) for resource in body['resources']]}, **filters}
                
            # obtenemos los recursos
            resources = list(mongodb.get_all_records('resources', filters, fields={'_id': 1}))
            resources = [str(resource['_id']) for resource in resources]

            records_filters = {
                'parent.id': {'$in': resources}
            }
        else:
            records_filters = {'_id': {'$in': [ObjectId(record) for record in body['records']]}}

        if 'overwrite' in body and body['overwrite']:
            records_filters = {"$or": [{"processing.liquidText": {"$exists": False}, **records_filters}, {"processing.liquidText": {"$exists": True}, **records_filters}]}
        else:
            records_filters['processing.liquidText'] = {'$exists': False}
        
        records = list(mongodb.get_all_records('records', records_filters, fields={'_id': 1, 'mime': 1, 'filepath': 1, 'processing': 1}))
        
        for r in records:
            text_key = None
            for k in r['processing']:
                if r['processing'][k]['type'] and r['processing'][k]['type'] == 'av_transcribe':
                    text_key = k
                    break
                
            if text_key:
                text = r['processing'][text_key]['result']['text']
                update = {
                    'processing': r['processing']
                }
                
                update['processing']['liquidText'] = {
                    'type': 'liquidText',
                    'result': {
                        'text': text,
                        'status': 'completed',
                        'date': datetime.datetime.now()
                    }
                }
                
                instance.update_data('records', str(r['_id']), update)
        
        return 'ok'
    
plugin_info = {
    'name': 'Generar texto líquido',
    'description': 'Generar texto líquido a partir de los archivos catalogados en el sistema',
    'version': '0.1',
    'author': '',
    'type': [],
    'settings': {
        'settings_bulk': []
    },
    'actions': [
        {
            'placement': 'detail_record',
            'record_type': ['audio', 'video'],
            'label': 'Crear texto líquido',
            'roles': ['admin', 'processing', 'editor'],
            'endpoint': 'bulk',
            'icon': 'WaterDrop, Article',
            'extraOpts': [
                {
                    'type': 'checkbox',
                    'label': 'Sobreescribir procesamientos existentes',
                    'id': 'overwrite',
                    'default': False,
                    'required': False,
                }
            ]
        },
        {
            'placement': 'detail_record',
            'record_type': ['audio', 'video'],
            'label': 'Descargar texto líquido',
            'roles': ['admin', 'processing', 'editor'],
            'endpoint': 'download',
            'icon': 'Download, WaterDrop',
            'extraOpts': [
                {
                    'type': 'select',
                    'label': 'Formato del archivo',
                    'id': 'format',
                    'default': 'pdf',
                    'options': [
                        {'value': 'pdf', 'label': 'PDF'},
                        {'value': 'doc', 'label': 'DOC'},
                        {'value': 'txt', 'label': 'TXT'},
                    ],
                    'required': False,
                }
            ]
        },
    ]
}