"""liquidText — turn a transcription into an editable, exportable text.

WHAT IT DOES. An audio or video record that has been transcribed carries its
transcript under ``processing``. This plugin copies that text into a
``liquidText`` processing entry that a person can then edit in the viewer, and
exports the result as DOCX, PDF or plain text.

THE CROSS-PLUGIN DEPENDENCY IS BEHIND AN INTERFACE. PDF export needs LibreOffice,
which lives in ``filesProcessing``. The conversion is looked up through
``archihub.plugins.framework.interop``, which answers a clear refusal when no
plugin provides it. Importing across plugin package boundaries instead couples
the two to each other's file layout, bypasses activation entirely — the import
succeeds whether or not the providing plugin is switched on — and reports its
failure as a sentence about a Python import from inside a Celery task.
"""

from __future__ import annotations

import datetime
import logging
import os

from celery import shared_task
from fastapi import Body, Depends
from fastapi.responses import JSONResponse

from archihub.core.i18n import gettext as _
from archihub.core.responses import json_response
from archihub.core.security.jwt import CurrentUser
from archihub.plugins.framework import data as plugin_data
from archihub.plugins.framework.base import (
    ArchiPlugin,
    BrokerUnavailable,
    as_response,
    object_ids,
    queue,
    require_roles,
    task_result_file,
)

logger = logging.getLogger(__name__)

SLUG = "liquidText"

TASK_GENERATE = "liquidText.generateLiquidText"
TASK_DOWNLOAD = "liquidText.downloadLiquidText"

#: Export formats. An ALLOWLIST indexed by the request, never a value pasted
#: into a filename: `format` reaches an extension and a code path.
FORMATS = ("doc", "pdf", "txt")

PROCESSING_KEY = "liquidText"

#: The processing entry kinds whose text this plugin can start from.
TRANSCRIPT_TYPES = ("av_transcribe",)


class LiquidText(ArchiPlugin):
    def add_routes(self) -> None:
        plugin = self

        @self.router.post(
            "/bulk",
            status_code=201,
            responses={201: {"description": "Queued"}, 503: {"description": "Queue unavailable"}},
        )
        def generate(
            body: dict = Body(...),
            current_user: CurrentUser = Depends(require_roles("admin", "processing")),
        ) -> JSONResponse:
            """Queue generation of liquid text for a set of records."""
            try:
                queue(
                    generate_task,
                    TASK_GENERATE,
                    current_user.username,
                    "msg",
                    body,
                    current_user.username,
                )
            except BrokerUnavailable:
                return json_response({"msg": _("The task queue is unavailable")}, 503)

            return json_response({"msg": _("The task was added to the processing queue")}, 201)

        @self.router.post(
            "/download",
            status_code=201,
            responses={201: {"description": "Queued"}, 400: {"description": "Unsupported format"}},
        )
        def download(
            body: dict = Body(...),
            current_user: CurrentUser = Depends(
                require_roles("admin", "processing", "editor", "transcriber")
            ),
        ) -> JSONResponse:
            """Queue an export of one record's liquid text.

            The format is checked HERE, not in the task. The legacy version
            queued whatever arrived and raised "Formato no soportado" inside the
            worker, so a typo came back as a failed background job minutes later
            instead of a 400.
            """
            if body.get("format") not in FORMATS:
                return json_response({"msg": _("Unsupported format")}, 400)

            try:
                queue(
                    download_task,
                    TASK_DOWNLOAD,
                    current_user.username,
                    "file_download",
                    body,
                    current_user.username,
                )
            except BrokerUnavailable:
                return json_response({"msg": _("The task queue is unavailable")}, 503)

            return json_response({"msg": _("The task was added to the processing queue")}, 201)

        @self.router.get(
            "/filedownload/{task_id}",
            responses={200: {"description": "The generated file"}, 404: {"description": "Unknown task"}},
        )
        def file_download(
            task_id: str,
            current_user: CurrentUser = Depends(
                require_roles("admin", "processing", "editor")
            ),
        ):
            """Download a completed export."""
            from archihub.api.users.services import has_role
            from archihub.core.responses import file_response

            result, status_code = task_result_file(
                task_id, current_user.username, is_admin=has_role(current_user.username, "admin")
            )
            if status_code != 200:
                return json_response(result, status_code)

            return file_response(
                result,
                download_name=os.path.basename(str(result)),
                as_attachment=True,
            )

        @self.router.post(
            "/save",
            status_code=201,
            responses={201: {"description": "Saved"}, 404: {"description": "Unknown record"}},
        )
        def save(
            body: dict = Body(...),
            current_user: CurrentUser = Depends(
                require_roles("admin", "processing", "editor", "transcriber")
            ),
        ) -> JSONResponse:
            """Save edited liquid text back onto a record."""
            return as_response(plugin.save_text(body, current_user.username))

    def save_text(self, body: dict, user: str) -> tuple[dict, int]:
        """Store an edited transcript against a record.

        THE RECORD'S OWN VISIBILITY IS CHECKED, not just the caller's global
        role. Checking the role alone lets an editor rewrite the transcript of a
        series they are not permitted to open — the write path must compose the
        same visibility rule the read path applies.
        """
        from archihub.api.records import services as record_services

        record_id = body.get("id")
        if not record_id:
            return {"msg": _("You must specify a {field}", field="id")}, 400

        text = body.get("text")
        if not isinstance(text, str):
            return {"msg": _("You must specify a {field}", field="text")}, 400

        record, error = record_services.load_visible(str(record_id), user)
        if error:
            return error

        stored = record.get("processing") or {}
        source = body.get("slug")
        if source and source not in stored:
            return {"msg": _("Record has not been processed with {slug}", slug=source)}, 404

        plugin_data.store_processing_result(
            str(record["_id"]),
            PROCESSING_KEY,
            {
                "type": PROCESSING_KEY,
                "result": {
                    "text": text,
                    "status": "completed",
                    "date": datetime.datetime.now(),
                },
            },
        )

        name = record.get("displayName") or record.get("name") or str(record["_id"])
        return {"msg": _("Liquid text saved for {resource}", resource=name)}, 201


# ---------------------------------------------------------------------------
# Tasks
# ---------------------------------------------------------------------------


def _mongo():
    from archihub.infra.mongo import get_mongo

    return get_mongo()


@shared_task(ignore_result=False, name=TASK_GENERATE)
def generate_task(body: dict, user: str) -> str:
    """Copy each matching record's transcript into a liquid-text entry."""
    filters = _record_filters(body)

    if body.get("overwrite"):
        pass  # every matching record, whether or not it already has one
    else:
        filters["processing.liquidText"] = {"$exists": False}

    records = list(
        _mongo().get_all_records(
            "records", filters, fields={"_id": 1, "processing": 1}
        )
    )

    written = 0
    for record in records:
        text = _transcript_of(record)
        if text is None:
            continue
        plugin_data.store_processing_result(
            str(record["_id"]),
            PROCESSING_KEY,
            {
                "type": PROCESSING_KEY,
                "result": {
                    "text": text,
                    "status": "completed",
                    "date": datetime.datetime.now(),
                },
            },
        )
        written += 1

    logger.info("liquidText: wrote %d of %d candidate records", written, len(records))
    return _("Liquid text generated for {count} records", count=written)


@shared_task(ignore_result=False, name=TASK_DOWNLOAD)
def download_task(body: dict, user: str) -> str:
    """Export one record's liquid text, returning the path relative to the user root."""
    from archihub.core import files as filestore
    from archihub.core.settings import get_settings

    fmt = body.get("format")
    if fmt not in FORMATS:
        raise ValueError(f"Unsupported format: {fmt!r}")

    record_ids = object_ids(body.get("records") or [], "records")
    if len(record_ids) != 1:
        # The original raised two different Spanish sentences for "none" and
        # "more than one"; both mean the same thing to the caller.
        raise ValueError("Select exactly one record to export")

    record = _mongo().get_record(
        "records",
        {"_id": record_ids[0]},
        fields={"_id": 1, "processing": 1, "name": 1, "displayName": 1},
    )
    if not record:
        raise ValueError("Record not found")

    entry = (record.get("processing") or {}).get(PROCESSING_KEY) or {}
    text = (entry.get("result") or {}).get("text")
    if not text:
        raise ValueError("This record has no liquid text to export")

    settings = get_settings()
    # `user` reaches a directory name. It is an authenticated username rather
    # than free text, but it is still data, and the original concatenated it.
    directory = filestore.resolve_within(settings.user_files_path, user, "liquidText")
    directory.mkdir(parents=True, exist_ok=True)

    title = record.get("displayName") or record.get("name") or str(record["_id"])
    stem = str(record["_id"])

    if fmt == "txt":
        destination = directory / f"{stem}.txt"
        destination.write_text(text, encoding="utf-8")
    elif fmt == "doc":
        destination = directory / f"{stem}.docx"
        write_docx(text, title, destination)
    else:
        destination = _export_pdf(text, title, stem, directory, settings)

    return f"/{user}/liquidText/{destination.name}"


def _export_pdf(text: str, title: str, stem: str, directory, settings):
    """DOCX first, then LibreOffice. Cleans up after itself either way."""
    from pathlib import Path

    from archihub.core import files as filestore
    from archihub.plugins.framework import interop

    scratch = filestore.resolve_within(settings.temporal_files_path, f"{stem}.docx")
    Path(scratch).parent.mkdir(parents=True, exist_ok=True)
    write_docx(text, title, scratch)

    destination = Path(directory) / f"{stem}.pdf"
    try:
        interop.convert_to_pdf(scratch, destination)
    finally:
        # The original removed the intermediate DOCX only on the success path,
        # so every failed export left one behind in the temporal volume.
        filestore.remove_quietly(scratch)

    if not destination.is_file():
        raise RuntimeError("PDF conversion produced no file")
    return destination


def write_docx(html: str, title: str, path) -> None:
    """Render the stored HTML fragment into a Word document.

    Only the inline emphasis the editor can produce is carried over; anything
    else becomes plain text. Kept as-is from the original, with one change: the
    original read the record's title from a loop variable that leaked in from
    the enclosing scope (`r['displayName']`), so refactoring the loop away would
    have raised `NameError`.
    """
    import re

    from bs4 import BeautifulSoup
    from docx import Document

    document = Document()
    document.add_heading(title, 0)

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6"]):
        if tag.name == "p":
            container = document.add_paragraph()
        elif re.fullmatch(r"h[1-6]", tag.name):
            container = document.add_heading(level=int(tag.name[1]))
        else:
            continue

        for node in tag.contents:
            name = getattr(node, "name", None)
            if name in ("b", "strong"):
                container.add_run(node.get_text()).bold = True
            elif name in ("i", "em"):
                container.add_run(node.get_text()).italic = True
            elif name == "u":
                container.add_run(node.get_text()).underline = True
            elif name is None:
                container.add_run(str(node))

    document.save(str(path))


def _transcript_of(record: dict) -> str | None:
    """The record's transcript text, if it has one."""
    for entry in (record.get("processing") or {}).values():
        if not isinstance(entry, dict):
            continue
        if entry.get("type") in TRANSCRIPT_TYPES:
            text = (entry.get("result") or {}).get("text")
            if isinstance(text, str) and text:
                return text
    return None


def _record_filters(body: dict) -> dict:
    """Which records a bulk run covers.

    Either an explicit list of record ids, or every record filed under the
    resources a content-type/parent selection matches.
    """
    if body.get("records"):
        return {"_id": {"$in": object_ids(body["records"], "records")}}

    post_type = body.get("post_type")
    if not post_type:
        raise ValueError(_("No content type was specified"))

    filters: dict = {"post_type": {"$in": post_type} if isinstance(post_type, list) else post_type}

    resources = body.get("resources") or []
    parent = body.get("parent")

    if parent and not resources:
        filters = {
            "$or": [
                {"parents.id": parent, "post_type": filters["post_type"]},
                {"_id": object_ids([parent], "parent")[0]},
            ]
        }
    elif resources:
        filters = {"_id": {"$in": object_ids(resources, "resources")}, **filters}

    matched = list(_mongo().get_all_records("resources", filters, fields={"_id": 1}))
    return {"parent.id": {"$in": [str(row["_id"]) for row in matched]}}


plugin_info = {
    "name": "Generar texto líquido",
    "description": "Generar texto líquido a partir de los archivos catalogados en el sistema",
    "version": "0.1",
    "author": "",
    "type": [],
    "settings": {"settings_bulk": []},
    "actions": [
        {
            "placement": "detail_record",
            "record_type": ["audio", "video"],
            "label": "Crear texto líquido",
            "roles": ["admin", "processing", "editor"],
            "endpoint": "bulk",
            "icon": "WaterDrop, Article",
            "extraOpts": [
                {
                    "type": "checkbox",
                    "label": "Sobreescribir procesamientos existentes",
                    "id": "overwrite",
                    "default": False,
                    "required": False,
                }
            ],
        },
        {
            "placement": "detail_record",
            "record_type": ["audio", "video"],
            "label": "Descargar texto líquido",
            "roles": ["admin", "processing", "editor"],
            "endpoint": "download",
            "icon": "Download, WaterDrop",
            "extraOpts": [
                {
                    "type": "select",
                    "label": "Formato del archivo",
                    "id": "format",
                    "default": "pdf",
                    "options": [
                        {"value": "pdf", "label": "PDF"},
                        {"value": "doc", "label": "DOC"},
                        {"value": "txt", "label": "TXT"},
                    ],
                    "required": False,
                }
            ],
        },
    ],
}


def build() -> LiquidText:
    return LiquidText(SLUG, plugin_info, module_file=__file__)
